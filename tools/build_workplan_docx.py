from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "KV_Cache_Privacy_Research_Workplan.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_keep(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_font(run, size=None, bold=None, color=INK, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering(doc: Document, bullet: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)
    run = p.add_run(text)
    set_font(run)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_repeat_keep(p)


def add_label_paragraph(doc: Document, label: str, text: str, after=4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r)


def add_role(doc: Document, title: str, question: str, responsibilities: list[str], gates: list[str], bullet_id: int) -> None:
    add_heading(doc, title, 2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Primary question. ")
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(question)
    set_font(r, italic=True)
    for item in responsibilities:
        add_list_item(doc, item, bullet_id)
    for gate in gates:
        add_label_paragraph(doc, "Acceptance gate", gate, after=5)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run("KV CACHE PRIVACY RESEARCH  |  INTERNAL WORKPLAN")
    set_font(hr, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Research workplan  |  ")
    set_font(fr, size=9, color=MUTED)
    add_page_number(footer)

    bullet_id = add_numbering(doc, bullet=True)
    number_id = add_numbering(doc, bullet=False)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(16)
    kicker.paragraph_format.space_after = Pt(3)
    kr = kicker.add_run("RESEARCH WORKPLAN")
    set_font(kr, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    tr = title.add_run("KV Cache Quantization x Privacy")
    set_font(tr, size=25, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run("Three-person execution plan for a controlled privacy-utility-bitrate study")
    set_font(sr, size=13.5, color=MUTED)

    add_label_paragraph(doc, "Status", "Kickoff plan - baseline-first research", after=2)
    add_label_paragraph(doc, "Initial model", "Llama-3.2-1B-Instruct; immutable revision to be pinned", after=2)
    add_label_paragraph(doc, "Initial precisions", "FP16, INT8, and INT4", after=2)
    add_label_paragraph(doc, "Deadline", "Monday, August 31, 2026", after=12)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    callout.style = "Table Grid"
    cell = callout.cell(0, 0)
    set_cell_shading(cell, PALE_BLUE)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cr = cp.add_run("Decision: ")
    set_font(cr, bold=True, color=NAVY)
    cr = cp.add_run("The first deliverable is a trustworthy privacy-utility-bitrate curve, not a new defense. Mathematical design begins only after the baseline reveals a stable phenomenon.")
    set_font(cr)

    add_heading(doc, "1. Research objective", 1)
    p = doc.add_paragraph("Determine how KV-cache precision changes exact, semantic, and sensitive-information leakage relative to model utility. Quantization is initially treated as a controlled independent variable, not as a privacy mechanism.")
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph("The core comparison is paired: the same prompt, checkpoint, seed, generation configuration, cache extraction point, and attack budget must be used at FP16, INT8, and INT4.")

    add_heading(doc, "2. Team allocation", 1)
    add_role(
        doc,
        "Person A - Quantization and utility",
        "Can the model use a lower-precision cache without losing task performance?",
        [
            "Pin the model and tokenizer revisions and deterministic generation configuration.",
            "Implement standardized K/V capture hooks and versioned cache bundles.",
            "Implement controlled symmetric INT8 and INT4 quantization with recorded scales, zero points, grouping, clipping, and RoPE stage.",
            "Integrate the quantized cache into autoregressive attention so utility is measured on a cache the model actually reads.",
            "Measure logit divergence, generation agreement, QA, long-context retrieval, memory, and latency.",
        ],
        [
            "Repeated FP16 extraction is tensor-identical and metadata-complete.",
            "One command reproduces the pilot online quantized-cache utility result.",
        ],
        bullet_id,
    )
    add_role(
        doc,
        "Person B - Leakage and attacks",
        "How much information can an informed attacker recover at each precision?",
        [
            "Reproduce a reliable FP16 reconstruction attack before interpreting quantization.",
            "Implement token, semantic, and sensitive-attribute leakage objectives.",
            "Evaluate both naive and quantization-adaptive attackers that know the quantizer and scales.",
            "Hold attack initialization, budget, stopping rule, and evaluation code constant across precisions.",
            "Write per-sample outputs with convergence diagnostics and later lead K-versus-V and layer/head analyses.",
        ],
        [
            "FP16 leakage is stable and exceeds a control baseline.",
            "Adaptive attacks consume quantized metadata correctly; a format mismatch cannot explain the result.",
        ],
        bullet_id,
    )
    add_role(
        doc,
        "Person C - Dataset, evaluation, and reproducibility",
        "Are the observed differences valid, statistically meaningful, and reproducible?",
        [
            "Create controlled synthetic public, personal, financial, medical, and confidential prompt categories.",
            "Predefine secret spans, attribute labels, immutable splits, and dataset versions.",
            "Maintain experiment manifests and validate paired comparisons across precisions.",
            "Aggregate per-sample leakage and utility records, confidence intervals, failure cases, and final figures.",
            "Audit cache schemas and later lead information-concentration, RAG, and multimodal extensions after the text baseline passes.",
        ],
        [
            "The dataset manifest, threat model, metrics, and cache schema are frozen before the final paired run.",
            "Every aggregate is traceable to per-sample records and a reproducible command.",
        ],
        bullet_id,
    )

    add_heading(doc, "3. Shared interface contract", 1)
    p = doc.add_paragraph("Person C publishes prompt manifests. Person A produces versioned cache bundles. Person B consumes those bundles and produces per-sample attack records. Person C joins leakage and utility outputs through prompt ID, precision, and experiment ID.")
    for item in [
        "Model/tokenizer revision, code revision, prompt ID, token IDs, and seed.",
        "Layer, KV-head, sequence, and head-dimension layout.",
        "Capture stage and whether RoPE has already been applied to keys.",
        "Original dtype, target precision, quantization granularity, scale, zero point, and clipping.",
        "Whether the quantized cache was used during inference or supplied only to the attacker.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "4. Threat model and attack levels", 1)
    add_label_paragraph(doc, "Attacker has", "Leaked KV cache, model architecture and weights, tokenizer, quantization algorithm, grouping, clipping rule, scales, and zero points.")
    add_label_paragraph(doc, "Attacker lacks", "Victim prompt, victim response, and secret labels.")
    for item in [
        "Naive attack - the FP16 attack is run without quantization-specific adaptation.",
        "Adaptive attack - the attacker dequantizes or explicitly optimizes against the quantized representation.",
        "Oracle diagnostic - optional upper-bound condition, always labeled unrealistic and never presented as an attack success.",
    ]:
        add_list_item(doc, item, number_id)

    warning = doc.add_table(rows=1, cols=1)
    set_table_geometry(warning, [9360])
    warning.style = "Table Grid"
    cell = warning.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    wp = cell.paragraphs[0]
    wr = wp.add_run("Claim boundary. ")
    set_font(wr, bold=True, color=NAVY)
    wr = wp.add_run("A drop in naive-attack performance is not evidence of privacy protection. Any security claim requires an adaptive attacker evaluated under the same model, prompts, precision, and budget.")
    set_font(wr)

    add_heading(doc, "5. Parallel execution", 1)
    doc.add_paragraph("All three people begin immediately and work concurrently. Person C publishes the smoke-test manifest and experiment identifiers; Person A and Person B independently produce cache/utility and attack outputs using those identifiers; Person C continuously validates compatibility and merges completed results.")
    for item in [
        "Person C publishes the shared smoke-test manifest before the final paired run.",
        "Person A and Person B do not wait for each other after the manifest and cache schema are fixed.",
        "A result enters the Monday report only when its reproduction command and limitations are recorded.",
    ]:
        add_list_item(doc, item, bullet_id)

    add_heading(doc, "6. Metrics and experimental distinction", 1)
    add_heading(doc, "Privacy pipeline", 2)
    doc.add_paragraph("Capture an FP16 cache once, quantize the same tensors at each precision, and supply those representations to the attacker.")
    add_heading(doc, "Utility pipeline", 2)
    doc.add_paragraph("Run autoregressive inference while attention actually reads the quantized cache. Post-hoc quantization of an unused cache cannot measure utility degradation.")

    metrics = doc.add_table(rows=1, cols=3)
    metrics.style = "Table Grid"
    for idx, text in enumerate(("Dimension", "Required measures", "Why separate")):
        cell = metrics.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        r = cell.paragraphs[0].add_run(text)
        set_font(r, bold=True, color=NAVY)
    metric_rows = [
        ("Exact leakage", "Token accuracy, sequence accuracy, edit distance", "Tests literal recovery"),
        ("Semantic leakage", "Embedding similarity plus inspectable examples", "Captures meaning after wording changes"),
        ("Sensitive leakage", "Attribute classification and secret-span recovery", "Measures the privacy harm directly"),
        ("Utility", "Logit divergence, generation agreement, QA, retrieval", "Separates privacy changes from model damage"),
        ("Efficiency", "Cache bytes, latency, throughput", "Confirms a real compression benefit"),
    ]
    for data in metric_rows:
        cells = metrics.add_row().cells
        for idx, value in enumerate(data):
            r = cells[idx].paragraphs[0].add_run(value)
            set_font(r, size=10)
    set_table_geometry(metrics, [1800, 4380, 3180])

    add_heading(doc, "7. Review gates", 1)
    gates = [
        "Protocol gate - no experiments before the model revision, threat model, cache schema, and dataset version are frozen.",
        "FP16 gate - quantization results are not interpreted until the FP16 attack is reliable.",
        "Adaptive-attack gate - naive degradation cannot be called a privacy benefit.",
        "Utility gate - utility requires an actually used quantized cache.",
        "Claim gate - confidence intervals, category results, and failure cases precede conclusions.",
    ]
    for gate in gates:
        add_list_item(doc, gate, number_id)

    add_heading(doc, "8. Repository workflow", 1)
    add_label_paragraph(doc, "Main branch", "Reviewed integration only.")
    add_label_paragraph(doc, "Person A branches", "feat/quantization-*")
    add_label_paragraph(doc, "Person B branches", "feat/attacks-*")
    add_label_paragraph(doc, "Person C branches", "feat/evaluation-*")
    p = doc.add_paragraph("Every experiment-changing pull request records the pinned revisions, seed, prompt-set version, cache schema, quantization parameters, inference/attack distinction, reproduction command, and verification performed.")

    add_heading(doc, "9. Monday submission package", 1)
    for item in [
        "One end-to-end FP16 smoke test.",
        "Controlled INT8 and INT4 conversion of the same caches.",
        "At least one adaptive leakage attack across the available precisions.",
        "An online quantized-cache utility smoke test, or an explicit blocker if integration is incomplete.",
        "Per-sample machine-readable results and initial privacy/utility versus precision plots.",
        "A concise limitations section separating completed results from planned work.",
    ]:
        add_list_item(doc, item, bullet_id)
    p = doc.add_paragraph("This is a pilot submission. It does not need to prove a new defense by Monday.")
    p.runs[0].bold = True

    add_heading(doc, "10. Later mathematical direction", 1)
    p = doc.add_paragraph("Only after the pilot reveals a stable phenomenon should the group optimize a privacy-aware quantizer. A candidate constrained objective is:")
    formula = doc.add_table(rows=1, cols=1)
    formula.style = "Table Grid"
    set_table_geometry(formula, [9360])
    set_cell_shading(formula.cell(0, 0), PALE_BLUE)
    fp = formula.cell(0, 0).paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("min_Q  L_utility(Q) + lambda_c R(Q) - lambda_p L_adaptive_attack(Q)")
    set_font(fr, size=11, bold=True, color=NAVY, name="Consolas")
    doc.add_paragraph("Here R(Q) is bitrate and the privacy term is evaluated against a strong adaptive attacker. Any proposed method must be compared with simple uniform, per-channel, and per-head quantization at matched bitrate and utility.")

    add_heading(doc, "11. Immediate kickoff checklist", 1)
    for item in [
        "Assign names to Persons A, B, and C.",
        "Pin the model and tokenizer revisions.",
        "Approve the synthetic-data policy and pilot size.",
        "Finalize the cache bundle schema.",
        "Select the first FP16 reconstruction attack.",
        "Create one end-to-end smoke test using five prompts.",
        "Hold the first integration review after the smoke test passes.",
    ]:
        add_list_item(doc, item, bullet_id)

    doc.core_properties.title = "KV Cache Quantization x Privacy - Three-Person Research Workplan"
    doc.core_properties.subject = "Three-person KV-cache privacy and utility research plan due August 31, 2026"
    doc.core_properties.author = "KV Cache Research Group"
    doc.core_properties.keywords = "KV cache, quantization, privacy, leakage, utility"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
